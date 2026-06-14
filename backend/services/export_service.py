from docx import Document
from openpyxl import Workbook
from reportlab.pdfgen import canvas
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from textwrap import wrap

def export_to_docx(
        question,
        answer
):

    document = Document()

    document.add_heading(
        "AI Assistant Response",
        0
    )

    document.add_paragraph(
        f"Question: {question}"
    )

    document.add_paragraph(
        f"Answer: {answer}"
    )

    file_name = (
        "exports/answer.docx"
    )

    document.save(
        file_name
    )

    return file_name



def export_to_excel(
        question,
        answer
):

    wb = Workbook()

    ws = wb.active

    ws.append(
        ["Question", "Answer"]
    )

    ws.append(
        [question, answer]
    )

    file_name = (
        "exports/answer.xlsx"
    )

    wb.save(
        file_name
    )

    return file_name

from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from textwrap import wrap

def export_to_pdf(
        question,
        answer
):

    file_name = "exports/answer.pdf"

    c = canvas.Canvas(
        file_name,
        pagesize=letter
    )

    width, height = letter

    y = height - 50

    c.setFont("Helvetica-Bold", 12)
    c.drawString(
        50,
        y,
        f"Question: {question}"
    )

    y -= 30

    c.setFont("Helvetica", 11)

    lines = wrap(
        f"Answer: {answer}",
        width=90
    )

    for line in lines:

        c.drawString(
            50,
            y,
            line
        )

        y -= 18

        if y < 50:
            c.showPage()
            c.setFont(
                "Helvetica",
                11
            )
            y = height - 50

    c.save()

    return file_name